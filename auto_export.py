#!/usr/bin/env python3
"""
Auto-Export New Articles to SSD

This script:
1. Monitors the news feed API for new articles
2. Exports new articles to Word/PDF documents
3. Optionally scrapes full article content
4. Saves everything to SSD folder
"""

import asyncio
import json
import os
import sys
import time
from datetime import datetime, timedelta
from pathlib import Path
from typing import List, Dict, Any, Optional, Set

import requests
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

# Configuration
CONFIG = {
    "api_url": "http://127.0.0.1:8000",
    "output_dir": "/Volumes/MustafaSSD/newsfeed_articles",
    "check_interval_minutes": 15,
    "export_format": "both",  # "word", "pdf", or "both"
    "scrape_full_content": True,
    "max_scrape_per_run": 50,
}


class ArticleTracker:
    """Track which articles have been exported."""

    def __init__(self, output_dir: str):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.tracker_file = self.output_dir / "exported_articles.json"
        self.exported_urls: Set[str] = set()
        self._load()

    def _load(self):
        """Load exported article URLs."""
        if self.tracker_file.exists():
            try:
                data = json.loads(self.tracker_file.read_text())
                self.exported_urls = set(data.get("urls", []))
            except:
                self.exported_urls = set()

    def _save(self):
        """Save exported article URLs."""
        data = {
            "urls": list(self.exported_urls),
            "last_updated": datetime.now().isoformat(),
            "total_exported": len(self.exported_urls),
        }
        self.tracker_file.write_text(json.dumps(data, indent=2))

    def is_exported(self, url: str) -> bool:
        """Check if article has been exported."""
        return url in self.exported_urls

    def mark_exported(self, urls: List[str]):
        """Mark articles as exported."""
        self.exported_urls.update(urls)
        self._save()

    def get_new_articles(self, articles: List[Dict]) -> List[Dict]:
        """Filter to only new articles."""
        return [a for a in articles if not self.is_exported(a.get("link", ""))]


def fetch_articles(limit: int = 1000) -> List[Dict[str, Any]]:
    """Fetch articles from the API."""
    try:
        response = requests.get(
            f"{CONFIG['api_url']}/articles",
            params={"limit": limit},
            timeout=300
        )
        if response.status_code == 200:
            data = response.json()
            return data.get("articles", [])
    except Exception as e:
        print(f"Error fetching articles: {e}")
    return []


def create_word_document(articles: List[Dict], output_path: Path) -> int:
    """Create Word document from articles."""
    doc = Document()

    # Title
    title = doc.add_heading("Medical News Feed - New Articles", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n").italic = True
    meta.add_run(f"New Articles: {len(articles)}").italic = True

    doc.add_paragraph()

    # Articles
    for i, article in enumerate(articles, 1):
        # Title
        heading = doc.add_heading(level=2)
        heading.add_run(f"{i}. {article.get('title', 'Untitled')}")

        # Source and date
        source_para = doc.add_paragraph()
        source = article.get('site', article.get('source', 'Unknown'))
        source_para.add_run(f"Source: {source}").bold = True
        if article.get('published'):
            source_para.add_run(f" | {article.get('published')}")

        # Summary
        if article.get('summary'):
            doc.add_paragraph(article['summary'])

        # Full content if available
        if article.get('full_text'):
            content_para = doc.add_paragraph()
            content_para.add_run("Full Article:\n").bold = True
            content_para.add_run(article['full_text'][:2000])
            if len(article.get('full_text', '')) > 2000:
                content_para.add_run("... [truncated]")

        # Link
        link_para = doc.add_paragraph()
        link_para.add_run("Link: ").bold = True
        link_para.add_run(article.get('link', 'N/A'))

        doc.add_paragraph()  # Spacer

    doc.save(str(output_path))
    return len(articles)


def create_pdf_document(articles: List[Dict], output_path: Path) -> int:
    """Create PDF document from articles."""
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        topMargin=0.5*inch,
        bottomMargin=0.5*inch
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=20,
        alignment=1,
        spaceAfter=12,
    )
    heading_style = ParagraphStyle(
        'ArticleHeading',
        parent=styles['Heading2'],
        fontSize=12,
        textColor=colors.HexColor('#0b6efd'),
        spaceAfter=6,
    )
    meta_style = ParagraphStyle(
        'Meta',
        parent=styles['Normal'],
        fontSize=9,
        textColor=colors.grey,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=6,
    )

    story = []

    # Title
    story.append(Paragraph("Medical News Feed - New Articles", title_style))
    story.append(Paragraph(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | {len(articles)} articles",
        meta_style
    ))
    story.append(Spacer(1, 20))

    # Articles
    for i, article in enumerate(articles, 1):
        title_text = f"{i}. {article.get('title', 'Untitled')}"
        # Escape HTML entities
        title_text = title_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        story.append(Paragraph(title_text, heading_style))

        source = article.get('site', article.get('source', 'Unknown'))
        published = article.get('published', '')
        source_text = f"<b>Source:</b> {source}"
        if published:
            source_text += f" | {published}"
        story.append(Paragraph(source_text, meta_style))

        if article.get('summary'):
            summary = article['summary'].replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            story.append(Paragraph(summary, body_style))

        story.append(Spacer(1, 10))

    doc.build(story)
    return len(articles)


async def scrape_full_content(articles: List[Dict]) -> List[Dict]:
    """Scrape full content for articles."""
    try:
        from full_article_scraper import HumanLikeScraper
    except ImportError:
        print("Full article scraper not available")
        return articles

    urls_to_scrape = [
        {"url": a.get("link", ""), "title": a.get("title", ""), "site": a.get("site", "")}
        for a in articles
        if a.get("link")
    ][:CONFIG["max_scrape_per_run"]]

    if not urls_to_scrape:
        return articles

    print(f"Scraping full content for {len(urls_to_scrape)} articles...")

    async with HumanLikeScraper(min_delay=3.0, max_delay=6.0, max_concurrent=2) as scraper:
        scraped_content = {}

        for i, url_info in enumerate(urls_to_scrape):
            url = url_info["url"]
            print(f"  [{i+1}/{len(urls_to_scrape)}] {url[:50]}...")

            content = await scraper.scrape_article(url)
            if content:
                scraped_content[url] = content.full_text

    # Merge scraped content back into articles
    for article in articles:
        url = article.get("link", "")
        if url in scraped_content:
            article["full_text"] = scraped_content[url]

    return articles


def export_new_articles(
    articles: List[Dict],
    output_dir: Path,
    export_format: str = "both"
) -> Dict[str, str]:
    """Export articles to files."""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    results = {}

    if export_format in ("word", "both"):
        word_path = output_dir / f"new_articles_{timestamp}.docx"
        count = create_word_document(articles, word_path)
        results["word"] = str(word_path)
        print(f"Saved {count} articles to Word: {word_path}")

    if export_format in ("pdf", "both"):
        pdf_path = output_dir / f"new_articles_{timestamp}.pdf"
        count = create_pdf_document(articles, pdf_path)
        results["pdf"] = str(pdf_path)
        print(f"Saved {count} articles to PDF: {pdf_path}")

    return results


async def run_export_cycle(tracker: ArticleTracker):
    """Run a single export cycle."""
    print(f"\n{'='*60}")
    print(f"Export cycle started at {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"{'='*60}")

    # Fetch articles
    print("Fetching articles from API...")
    articles = fetch_articles(limit=2000)
    print(f"Found {len(articles)} total articles")

    # Filter to new articles only
    new_articles = tracker.get_new_articles(articles)
    print(f"New articles to export: {len(new_articles)}")

    if not new_articles:
        print("No new articles to export.")
        return

    # Optionally scrape full content
    if CONFIG["scrape_full_content"]:
        new_articles = await scrape_full_content(new_articles)

    # Export to files
    output_dir = Path(CONFIG["output_dir"])
    output_dir.mkdir(parents=True, exist_ok=True)

    results = export_new_articles(
        new_articles,
        output_dir,
        CONFIG["export_format"]
    )

    # Mark as exported
    exported_urls = [a.get("link", "") for a in new_articles if a.get("link")]
    tracker.mark_exported(exported_urls)

    print(f"\nExport complete!")
    print(f"Total exported this session: {len(new_articles)}")
    print(f"Total tracked articles: {len(tracker.exported_urls)}")


async def run_continuous():
    """Run continuous monitoring and export."""
    print("="*60)
    print("Medical News Feed Auto-Exporter")
    print("="*60)
    print(f"Output directory: {CONFIG['output_dir']}")
    print(f"Check interval: {CONFIG['check_interval_minutes']} minutes")
    print(f"Export format: {CONFIG['export_format']}")
    print(f"Full content scraping: {CONFIG['scrape_full_content']}")
    print("="*60)

    tracker = ArticleTracker(CONFIG["output_dir"])
    print(f"Previously exported: {len(tracker.exported_urls)} articles")

    while True:
        try:
            await run_export_cycle(tracker)
        except Exception as e:
            print(f"Error in export cycle: {e}")

        # Wait for next cycle
        print(f"\nNext check in {CONFIG['check_interval_minutes']} minutes...")
        await asyncio.sleep(CONFIG["check_interval_minutes"] * 60)


async def run_once():
    """Run a single export cycle."""
    tracker = ArticleTracker(CONFIG["output_dir"])
    await run_export_cycle(tracker)


def main():
    """Main entry point."""
    import argparse

    parser = argparse.ArgumentParser(description="Auto-export new medical news articles")
    parser.add_argument("--once", action="store_true", help="Run once and exit")
    parser.add_argument("--output", type=str, default=CONFIG["output_dir"],
                       help="Output directory")
    parser.add_argument("--format", choices=["word", "pdf", "both"], default="both",
                       help="Export format")
    parser.add_argument("--no-scrape", action="store_true",
                       help="Disable full content scraping")
    parser.add_argument("--interval", type=int, default=15,
                       help="Check interval in minutes")

    args = parser.parse_args()

    CONFIG["output_dir"] = args.output
    CONFIG["export_format"] = args.format
    CONFIG["scrape_full_content"] = not args.no_scrape
    CONFIG["check_interval_minutes"] = args.interval

    if args.once:
        asyncio.run(run_once())
    else:
        asyncio.run(run_continuous())


if __name__ == "__main__":
    main()
