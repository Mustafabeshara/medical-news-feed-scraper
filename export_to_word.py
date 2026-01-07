#!/usr/bin/env python3
"""Export all articles to Word document."""

import requests
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from collections import defaultdict
import sys

def fetch_all_articles():
    """Fetch all articles using requests in batches."""
    all_articles = []
    offset = 0
    batch_size = 5000  # Larger batch since we have proper offset now

    print("Fetching articles...")
    sys.stdout.flush()

    session = requests.Session()

    # First get total count
    try:
        response = session.get("http://127.0.0.1:8000/articles?limit=1&offset=0", timeout=30)
        data = response.json()
        total = data.get("total", 0)
        print(f"Total articles available: {total}")
    except Exception as e:
        print(f"Error getting total: {e}")
        total = 100000  # fallback

    while offset < total:
        try:
            url = f"http://127.0.0.1:8000/articles?limit={batch_size}&offset={offset}"
            response = session.get(url, timeout=120)

            if response.status_code != 200:
                print(f"HTTP {response.status_code} at offset {offset}")
                break

            data = response.json()
            articles = data.get("articles", [])

            if not articles:
                print(f"No more articles at offset {offset}")
                break

            all_articles.extend(articles)
            print(f"Batch offset {offset}: {len(articles)} articles (total: {len(all_articles)})")
            sys.stdout.flush()
            offset += batch_size

            if len(articles) < batch_size:
                print(f"Last batch (only {len(articles)} articles)")
                break
        except Exception as e:
            print(f"Error at offset {offset}: {e}")
            break

    return all_articles

def create_word_document(articles, output_path):
    """Create Word document from articles."""
    doc = Document()

    # Title
    title = doc.add_heading("Medical News Feed - Complete Article Export", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Export Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Total Articles: {len(articles)}")
    doc.add_paragraph("")

    # Group by source
    by_source = defaultdict(list)
    for art in articles:
        source = art.get("source", "Unknown")
        by_source[source].append(art)

    print(f"Grouped into {len(by_source)} sources")
    sys.stdout.flush()

    # Add articles
    count = 0
    for source in sorted(by_source.keys()):
        arts = by_source[source]
        doc.add_heading(f"{source} ({len(arts)} articles)", level=1)

        for art in arts:
            # Title
            title_text = art.get("title", "No Title")
            title_text = title_text.replace("<![CDATA[", "").replace("]]>", "")
            p = doc.add_paragraph()
            run = p.add_run(title_text)
            run.bold = True
            run.font.size = Pt(11)

            # Date
            pub = art.get("published", "")[:16] if art.get("published") else "No date"
            doc.add_paragraph(f"Date: {pub}")

            # Link
            link = art.get("link", "")
            if link:
                doc.add_paragraph(f"Link: {link}")

            # Companies/Products
            companies = art.get("companies", [])
            products = art.get("products", [])
            if companies:
                doc.add_paragraph(f"Companies: {', '.join(companies)}")
            if products:
                doc.add_paragraph(f"Products: {', '.join(products)}")

            # Summary
            summary = art.get("summary", "")
            summary = summary.replace("<![CDATA[", "").replace("]]>", "")
            if summary:
                doc.add_paragraph(f"Summary: {summary[:400]}...")

            doc.add_paragraph("")  # Spacing
            count += 1

            if count % 5000 == 0:
                print(f"Processed {count} articles...")
                sys.stdout.flush()

    doc.save(output_path)
    return count

def main():
    print("Starting export...")
    sys.stdout.flush()

    # Fetch articles
    articles = fetch_all_articles()
    print(f"\nTotal fetched: {len(articles)} articles")
    sys.stdout.flush()

    if not articles:
        print("No articles to export!")
        return

    # Create Word document
    output_path = "/Users/mustafaahmed/Desktop/Medical_News_Articles.docx"
    count = create_word_document(articles, output_path)

    print(f"\nSaved {count} articles to: {output_path}")

if __name__ == "__main__":
    main()
