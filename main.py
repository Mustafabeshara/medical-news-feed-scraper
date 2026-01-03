import asyncio
import io
import os
import time
from datetime import datetime
from typing import Dict, List, Any, Optional

import yaml
from fastapi import FastAPI, Query
from fastapi.responses import HTMLResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

from aggregator import fetch_all_sites, filter_articles
from config import CONFIG

REFRESH_INTERVAL_SEC = CONFIG.refresh_interval_sec
CONFIG_PATH = os.getenv("SITES_CONFIG", os.path.join(os.path.dirname(__file__), "sites.yaml"))

app = FastAPI(
    title="Medical News Feed Scraper",
    description="Aggregate medical news from multiple sources with smart feed discovery and scraping",
    version="2.0.0"
)

# In-memory cache
_cache_articles_by_site: Dict[str, List[Dict[str, Any]]] = {}
_cache_last_refresh: float = 0.0
_sites_config: List[Dict[str, Any]] = []


def load_config() -> List[Dict[str, Any]]:
    with open(CONFIG_PATH, "r", encoding="utf-8") as f:
        data = yaml.safe_load(f) or {}
    sites = data.get("sites") or []
    return sites


async def refresh() -> None:
    """Fetch articles from all sites using concurrent requests."""
    global _cache_articles_by_site, _cache_last_refresh

    print(f"Starting refresh of {len(_sites_config)} sites...")
    start_time = time.time()

    # Use new concurrent fetch_all_sites
    results = await fetch_all_sites(_sites_config)

    # Update cache with results
    _cache_articles_by_site.update(results)
    _cache_last_refresh = time.time()

    elapsed = time.time() - start_time
    total_articles = sum(len(articles) for articles in results.values())
    print(f"Refresh complete: {len(results)} sites, {total_articles} articles in {elapsed:.2f}s")


async def refresher_task() -> None:
    while True:
        await refresh()
        await asyncio.sleep(REFRESH_INTERVAL_SEC)


@app.on_event("startup")
async def on_startup():
    global _sites_config
    _sites_config = load_config()
    # Start background refresher (initial refresh happens immediately in the task)
    asyncio.create_task(refresher_task())


@app.get("/health")
async def health_check():
    """Health check endpoint showing system status."""
    return {
        "status": "healthy",
        "last_refresh": _cache_last_refresh,
        "last_refresh_iso": datetime.fromtimestamp(_cache_last_refresh).isoformat() if _cache_last_refresh > 0 else None,
        "sites_configured": len(_sites_config),
        "sites_with_articles": len(_cache_articles_by_site),
        "total_articles": sum(len(articles) for articles in _cache_articles_by_site.values()),
        "version": "2.0.0",
        "refresh_interval_sec": REFRESH_INTERVAL_SEC,
    }


@app.get("/sites")
async def get_sites() -> List[str]:
    """Get list of all sites with articles."""
    return list(_cache_articles_by_site.keys())


@app.get("/articles")
async def get_articles(
    site: Optional[str] = Query(None),
    q: Optional[str] = Query(None),
    limit: int = Query(50, ge=1, le=500),
):
    # flatten all articles
    all_articles: List[Dict[str, Any]] = []
    for _, items in _cache_articles_by_site.items():
        all_articles.extend(items)
    filtered = filter_articles(all_articles, site=site, q=q, limit=limit)
    return {
        "last_refresh": _cache_last_refresh,
        "count": len(filtered),
        "articles": filtered,
    }


@app.get("/", response_class=HTMLResponse)
async def index():
    # Serve minimal HTML from static
    with open(
        os.path.join(os.path.dirname(__file__), "static", "index.html"), "r", encoding="utf-8"
    ) as f:
        return HTMLResponse(f.read())


def _get_filtered_articles(site: Optional[str], q: Optional[str], limit: int) -> List[Dict[str, Any]]:
    """Helper to get filtered articles for export."""
    all_articles: List[Dict[str, Any]] = []
    for _, items in _cache_articles_by_site.items():
        all_articles.extend(items)
    return filter_articles(all_articles, site=site, q=q, limit=limit)


@app.get("/export/word")
async def export_word(
    site: Optional[str] = Query(None),
    q: Optional[str] = Query(None),
    limit: int = Query(50, ge=1, le=500),
):
    """Export articles to Word document."""
    articles = _get_filtered_articles(site, q, limit)

    doc = Document()

    # Title
    title = doc.add_heading("Medical News Feed", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n").italic = True
    meta.add_run(f"Total Articles: {len(articles)}").italic = True
    if site:
        meta.add_run(f" | Site: {site}").italic = True
    if q:
        meta.add_run(f" | Search: {q}").italic = True

    doc.add_paragraph()

    # Articles
    for i, article in enumerate(articles, 1):
        # Article title
        heading = doc.add_heading(level=2)
        heading.add_run(f"{i}. {article.get('title', 'Untitled')}")

        # Source and date
        source_para = doc.add_paragraph()
        source_para.add_run(f"Source: {article.get('site', article.get('source', 'Unknown'))}").bold = True
        if article.get('published'):
            source_para.add_run(f" | {article.get('published')}")

        # Summary
        if article.get('summary'):
            doc.add_paragraph(article['summary'])

        # Link
        link_para = doc.add_paragraph()
        link_para.add_run("Link: ").bold = True
        link_para.add_run(article.get('link', 'N/A'))

        doc.add_paragraph()  # Spacer

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"medical_news_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


@app.get("/export/pdf")
async def export_pdf(
    site: Optional[str] = Query(None),
    q: Optional[str] = Query(None),
    limit: int = Query(50, ge=1, le=500),
):
    """Export articles to PDF document."""
    articles = _get_filtered_articles(site, q, limit)

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        alignment=1,  # Center
        spaceAfter=12,
    )
    heading_style = ParagraphStyle(
        'ArticleHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#0b6efd'),
        spaceAfter=6,
    )
    meta_style = ParagraphStyle(
        'Meta',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.grey,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=6,
    )
    link_style = ParagraphStyle(
        'Link',
        parent=styles['Normal'],
        fontSize=9,
        textColor=colors.HexColor('#0066cc'),
        spaceAfter=12,
    )

    story = []

    # Title
    story.append(Paragraph("Medical News Feed", title_style))

    # Metadata
    meta_text = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | Total Articles: {len(articles)}"
    if site:
        meta_text += f" | Site: {site}"
    if q:
        meta_text += f" | Search: {q}"
    story.append(Paragraph(meta_text, meta_style))
    story.append(Spacer(1, 20))

    # Articles
    for i, article in enumerate(articles, 1):
        # Title
        title_text = f"{i}. {article.get('title', 'Untitled')}"
        story.append(Paragraph(title_text, heading_style))

        # Source
        source = article.get('site', article.get('source', 'Unknown'))
        published = article.get('published', '')
        source_text = f"<b>Source:</b> {source}"
        if published:
            source_text += f" | {published}"
        story.append(Paragraph(source_text, meta_style))

        # Summary
        if article.get('summary'):
            # Escape HTML entities in summary
            summary = article['summary'].replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            story.append(Paragraph(summary, body_style))

        # Link
        link = article.get('link', 'N/A')
        story.append(Paragraph(f"<b>Link:</b> {link}", link_style))
        story.append(Spacer(1, 10))

    doc.build(story)
    buffer.seek(0)

    filename = f"medical_news_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


# Mount static assets
app.mount(
    "/static",
    StaticFiles(directory=os.path.join(os.path.dirname(__file__), "static")),
    name="static",
)
